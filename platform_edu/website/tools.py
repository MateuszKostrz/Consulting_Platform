#!/usr/bin/env python
"""
Module Name: tools

Module provides a set of functions and classes to facilitate the training, testing, and evaluation of GPT models.
"""

# IMPORT LIBRARIES
from docx import Document
from typing import Tuple, List, Dict, Any, Optional, Union
import glob
import os
import re
import json
import numpy as np
from datetime import datetime
import matplotlib.pyplot as plt
import base64
import tiktoken
from pydantic import BaseModel
from functools import wraps
import logging
from pdf2docx import Converter
import io
from werkzeug.datastructures import FileStorage
import mimetypes
from io import TextIOWrapper
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.core.files.uploadedfile import TemporaryUploadedFile
from django.conf import settings

SYSTEM_PROMPT_S = '''

Use the following step-by-step instructions to respond to user input.
1. **Personal Statement Submission**: User will submit a Personal Statement enclosed in <ps> XML tags.

2. **Evaluation**: Evaluate the Personal Statement based on the Factors and instructions provided within the ***Scoring Guidelines***.
At the top of the Personal Statement student declares course She applies to. Content of the Personal Statement should present both student's motivation why She is a good fit both for course declared at the top.
For every Factor, in a step by step manner, determine whether conditions specified in ***Scoring Guidelines*** are fulfilled and assign a score accordingly. 
Take time to reason before answering.

    **Scoring Guidelines**:
    - Factor 1. Relevance and Motivation:
        - Grant 1-2 points if: The statement lacks clear motivation or fails to demonstrate any connection to the chosen course.
        - Grant 3-5 points if: The statement shows some general motivation and relevance, but examples are generic or not explicitly linked to the course.
        - Grant 6-8 points if: Motivation is articulated clearly, with relevant personal experiences and some alignment to the course, but lacks depth or consistent detail.
        - Grant 9-10 points if: Motivation is compelling and specifically tailored to the course, with personal experiences linked directly to subject choice, showing deep understanding of the field.
    - Factor 2. Experience and Achievements:
        - Grant 1-2 points if: The statement either does not mention any relevant experiences or achievements, or they are completely unrelated to the course.
        - Grant 3-5 points if: The statement includes some relevant experiences and achievements, but lacks consistent depth or clarity in explanation.
        - Grant 6-8 points if: The statement highlights relevant achievements and experiences, demonstrating suitability for the course, but with minor gaps in explanation or specificity.
        - Grant 9-10 points if: The statement provides a comprehensive overview of significant and specific achievements directly related to the course, showing preparation and enthusiasm for the field.
    - Factor 3. Skills and Personal Qualities:
        - Grant 1-2 points if: The statement does not identify relevant skills or qualities; no examples provided.
        - Grant 3-4 points if: The statement mentions relevant skills and qualities with examples, but explanations lack depth or clarity.
        - Grant 5 points if: The statement exemplifies key skills and personal qualities with strong, detailed examples, connecting them explicitly to the course requirements and real-world application.
    - Factor 4. Structure and Coherence:
        - Grant 1-2 points if: The statement is poorly structured, lacks logical flow, and is difficult to follow.
        - Grant 3-4 points if: The statement, generally, shows clear structure, but minor issues with flow or clarity; transitions may be weak.
        - Grant 5 points if: The statement is exceptionally well-organized with a clear introduction, thematic sections, and a compelling conclusion that reinforces the candidate’s suitability for the course.
    - Factor 5. Language and Presentation:
        - Grant 1-2 points if: Numerous grammar, spelling, or syntax errors significantly distract from the meaning, making the statement appear unprofessional.
        - Grant 3-4 points if: Some language errors are to be spotted, but overall meaning is clear; presentation is generally professional.
        - Grant 5 points if: The statement presents flawless language use, with varied sentence structure, precise vocabulary, and a polished presentation reflecting attention to detail and professionalism.

3. **Marking**: For each Factor assign a score reflecting how well the Personal Statement meets the factor and the conditions
specified in ***Scoring Guidelines***. 

4. **Output**: Provide the results in a JSON format as follows: {{"factor1":8, "factor2":5, ..., "factor5":4}}

Reason thorougly before answering. 
Evaluate every condition in ***Scoring Guidelines*** and every factor in a step by step manner and only afterwards answer in a specified JSON format.
'''

USER_PROMPT_S = '''
<ps>
{ps}
</ps>'''

SYSTEM_PROMPT_R = '''
Use the following step-by-step instructions to respond to user input.
1. **Personal Statement Submission**: User will submit
    - Personal Statement enclosed in <ps> XML tags.
    - A List of scores enclosed in `<scores>` XML tags.
- Scores indicate how well grading factors, listed within ***Scoring Guidelines*** section, are met. 
At the top of the Personal Statement student declares course She applies to. Content of the Personal Statement should present both student's motivation why She is a good fit both for course declared at the top.
***Scoring Guidelines*** section, apart from grading factors, comprises instructions on how to evaluate subsequent factors.

    **Scoring Guidelines**:
    - Factor 1. Relevance and Motivation:
        - Grant 1-2 points if: The statement lacks clear motivation or fails to demonstrate any connection to the chosen course.
        - Grant 3-5 points if: The statement shows some general motivation and relevance, but examples are generic or not explicitly linked to the course.
        - Grant 6-8 points if: Motivation is articulated clearly, with relevant personal experiences and some alignment to the course, but lacks depth or consistent detail.
        - Grant 9-10 points if: Motivation is compelling and specifically tailored to the course, with personal experiences linked directly to subject choice, showing deep understanding of the field.
    - Factor 2. Experience and Achievements:
        - Grant 1-2 points if: The statement either does not mention any relevant experiences or achievements, or they are completely unrelated to the course.
        - Grant 3-5 points if: The statement includes some relevant experiences and achievements, but lacks consistent depth or clarity in explanation.
        - Grant 6-8 points if: The statement highlights relevant achievements and experiences, demonstrating suitability for the course, but with minor gaps in explanation or specificity.
        - Grant 9-10 points if: The statement provides a comprehensive overview of significant and specific achievements directly related to the course, showing preparation and enthusiasm for the field.
    - Factor 3. Skills and Personal Qualities:
        - Grant 1-2 points if: The statement does not identify relevant skills or qualities; no examples provided.
        - Grant 3-4 points if: The statement mentions relevant skills and qualities with examples, but explanations lack depth or clarity.
        - Grant 5 points if: The statement exemplifies key skills and personal qualities with strong, detailed examples, connecting them explicitly to the course requirements and real-world application.
    - Factor 4. Structure and Coherence:
        - Grant 1-2 points if: The statement is poorly structured, lacks logical flow, and is difficult to follow.
        - Grant 3-4 points if: The statement, generally, shows clear structure, but minor issues with flow or clarity; transitions may be weak.
        - Grant 5 points if: The statement is exceptionally well-organized with a clear introduction, thematic sections, and a compelling conclusion that reinforces the candidate’s suitability for the course.
    - Factor 5. Language and Presentation:
        - Grant 1-2 points if: Numerous grammar, spelling, or syntax errors significantly distract from the meaning, making the statement appear unprofessional.
        - Grant 3-4 points if: Some language errors are to be spotted, but overall meaning is clear; presentation is generally professional.
        - Grant 5 points if: The statement presents flawless language use, with varied sentence structure, precise vocabulary, and a polished presentation reflecting attention to detail and professionalism.

2. **Evaluation**: For every factor:
   - **Interpret the Score**: Use the provided ***Scoring Guidelines*** to understand the extent to which the factor is met.
   - **Provide Reasoning**: Offer a thorough explanation of why the factor is met to the extent specified by the score. 
   Use the specific examples from the Personal Statement provided by the user to illustrate how the criterion is fulfilled. 
   Clearly indicate which parts of the Personal Statement correspond to the criterion.

3. **Output**: Present the reasonings in a JSON format as follows: {{"factor1":reasoning#1, "factor2":reasoning#2,...,"factor5":reasoning#5}}
'''

USER_PROMPT_R = '''
<ps>
{ps}
</ps>

**Scores**:
<scores>
{scores}
</scores>
'''

EXAMPLES_PATH = os.path.join(settings.BASE_DIR, 'website', 'data')

METRICS = ['Root Mean Squared Error', 'Mean Bias Error', 'Total Error Normalized']
TEST_LOGS_PATH = 'test_logs'
TEST_SCORES_PATH = 'test_scores'

# DEFINE CLASSES AND FUNCTIONS
class ScoresExtraction(BaseModel):
    factor1: int
    factor2: int
    factor3: int
    factor4: int
    factor5: int

class ReasoningsExtraction(BaseModel):
    factor1: str
    factor2: str
    factor3: str
    factor4: str
    factor5: str

class DocumentNotLoadedError(Exception):
    """
    Exception raised when a document is accessed before being loaded.
    """
    def __init__(self, message="The document has not been loaded. Please load a document before accessing the method."):
        """
        Initialize the DocumentNotLoadedError with an optional error message.
        
        Args:
            message (str): The error message to be displayed.
        """
        self.message = message
        super().__init__(self.message)


def get_image_description(client: Any, b64_image: str) -> str:
    """
    Generates a textual description of the content within an image using a GPT-4-based model.

    Args:
        client (Any): An instance of an API client used to interact with the GPT-4 service.
        b64_image (str): A base64-encoded string representing the image to be described.

    Returns:
        str: A concise description of the image content as generated by the model.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Describe what's in the image. Be precise and concise."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_image}"}}
                ]
            }
        ],
        max_tokens=300,
    )
    
    return response.choices[0].message.content

class TOKReader:
    """
    Custom class to read and process TOK IB Exhibitions.
    """

    def __init__(self):
        self.namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
    
    

    def load_doc(self, file: Union[str, FileStorage, TextIOWrapper, TemporaryUploadedFile]) -> None:
        """
        Loads a PDF or DOCX file into a `Document` object.

        Function handles loading both PDF and DOCX formats. It converts PDF files
        into DOCX format before loading them into a `Document` object. The file can be
        provided as a string path, a Flask `FileStorage` object, a Django `TemporaryUploadedFile`,
        or a `TextIOWrapper`.

        Args:
            file (str or FileStorage or TextIOWrapper or TemporaryUploadedFile): The file to load, either as a string
            representing the file path, or a file-like object.

        Raises:
            FileNotFoundError: If the file path is invalid or doesn't exist.
            ValueError: If the file format is unsupported (neither PDF nor DOCX).
            RuntimeError: If any other error occurs while loading the document.
        """
        try:
            print("Making it here")
            print(f"Type of file: {type(file)}")  # Debugging line

            # Handle different types of file inputs
            if isinstance(file, str):
                if not os.path.exists(file):
                    raise FileNotFoundError(f'File not found: {file}')
                _, file_ext = os.path.splitext(file)
                print(f"File extension from path: {file_ext}")  # Debugging line
                with open(file, 'rb') as f:
                    file_stream = io.BytesIO(f.read())
            elif isinstance(file, (InMemoryUploadedFile, TemporaryUploadedFile)):  # Updated to include TemporaryUploadedFile
                file_ext = os.path.splitext(file.name)[1]  # Get file extension
                print(f"File extension from Django uploaded file: {file_ext}")  # Debugging line
                file_stream = io.BytesIO(file.read())
            elif isinstance(file, TextIOWrapper):
                file_ext = os.path.splitext(file.name)[1] if hasattr(file, 'name') else ''
                print(f"File extension from TextIOWrapper: {file_ext}")  # Debugging line
                file_stream = io.BytesIO(file.read().encode('utf-8'))
            else:
                raise ValueError('Invalid file type. Must be a string path or file-like object.')

            print(f"File extension being processed: {file_ext}")  # Debugging line

            if file_ext.lower() == '.pdf':
                file_stream.seek(0)
                docx_buffer = io.BytesIO()
                converter = Converter(stream=file_stream)
                converter.convert(docx_buffer)
                docx_buffer.seek(0)
                self.doc = Document(docx_buffer)
            elif file_ext.lower() == '.docx':
                self.doc = Document(file_stream)
            else:
                raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")
        except Exception as e:
            print(f"[ERROR] Error processing uploaded file: {e}")  # Updated error message
            raise RuntimeError(f"An error occurred while loading the document: {e}")





    def load_body(self, client=None) -> str:
        """
        Extracts and returns the main content of the loaded document, with special handling for paragraphs
        and embedded images.

        Args:
            client (optional): An optional client instance for sending images to OpenAI for description.

        Returns:
            str: The extracted content of the document as a single string.

        Raises:
            DocumentNotLoadedError: If no document has been loaded into the instance.

        Notes:
            - Images, if in the document, are described using OpenAI and included in
            the content between `[IMAGE DESCRIPTION START]` and `[IMAGE DESCRIPTION END]` tags.
            - Content stops extracting when certain keywords like 'grade', 'grading criteria', or 'total points'
            are encountered.
        """

        if not hasattr(self, 'doc'):
            raise DocumentNotLoadedError()

        if hasattr(self, '_body_cache') and self._body_cache is not None:
            return self._body_cache

        paragraphs = []

        if not self.doc.element.body.findall('.//w:drawing', self.namespaces):
            for paragraph in self.doc.paragraphs:
                text = paragraph.text.strip()
                if text.lower().startswith(('grade', 'grading criteria', 'total points')):
                    break
                paragraphs.append(text)
            self._body_cache = " ".join(paragraphs)
            return self._body_cache

        for paragraph in self.doc.element.body.findall('.//w:p', self.namespaces):
            text = paragraph.text.strip()
            if text.lower().startswith('total:'):
                break
            paragraphs.append(text)

            drawings = paragraph.findall('.//w:drawing', self.namespaces)
            if drawings:
                for drawing in drawings:
                    paragraphs.append('\n[IMAGE DESCRIPTION START]\n')
                    blip = drawing.find('.//a:blip', namespaces=self.namespaces)
                    image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    image = self.doc.part.related_parts[image_id].blob
                    b64_image = base64.b64encode(image).decode('utf-8')

                    # Sends image to OpenAI.
                    image_description = get_image_description(client, b64_image)
                    paragraphs.append(image_description)
                    paragraphs.append('\n[IMAGE DESCRIPTION END]\n')
        return ' '.join(paragraphs)

    
    def load_grading(self) -> Tuple[List[str], List[str], List[str]]:
        """
        Extracts grading information from the loaded document's table and returns it as a tuple of lists.

        Returns:
            Tuple[List[str], List[str], List[str]]: A tuple containing three lists:
                - criteria: Grading criteria extracted from the table.
                - scales: Scales associated with the criteria.
                - reasonings: Reasonings or descriptions related to the criteria.
                - scores

        Raises:
            DocumentNotLoadedError: If no document has been loaded into the instance.
            ValueError: If no table is found in the document.

        Notes:
            - The method looks for tables (`w:tbl`) in the document's body to extract grading information.
            - Rows in the table must contain four valid cells with criteria, scale, reasoning, and a numeric score.
            Rows not meeting the format are skipped.
            - Results are cached to optimize subsequent calls.
        """

        if not hasattr(self, 'doc'):
            raise DocumentNotLoadedError()

        if hasattr(self, '_grading_cache') and self._grading_cache is not None:
            return self._grading_cache

        if not self.doc.element.body.findall('.//w:tbl', self.namespaces):
            raise ValueError('Table not found.')
            
        criteria = []
        reasonings = []
        scales = []
        scores = []

        for table in self.doc.element.body.findall('.//w:tbl', self.namespaces):
            for row in table.findall('w:tr', self.namespaces):
                criterion, scale, score, reasoning = [
                    ' '.join([p.text for p in cell.findall('w:p', self.namespaces)]).strip()
                    for cell in row.findall('w:tc', self.namespaces)
                ]
                if not (criterion and reasoning and scale and score and score.isdigit()):
                    continue
                criteria.append(criterion)
                scales.append(scale)
                reasonings.append(reasoning)
                scores.append(score)

        self._grading_cache = (criteria, scales, reasonings, scores)
        return self._grading_cache
        
    def _create_model_input_scores(self, client:Any=None) -> Tuple[str, str]:
        """
        Creates input prompts for a model by combining predefined system and user prompts
        with the document's body content.

        Args:
            client (Any, optional): An optional client instance for use in `load_body`. Defaults to None.

        Returns:
            Tuple[str, str]: A tuple containing:
                - system_prompt (str): The predefined system prompt.
                - user_prompt (str): The user prompt formatted with the extracted essay body.

        Notes:
            - The essay body is extracted using `load_body`, which may include image descriptions
            if applicable.
            - `SYSTEM_PROMPT_S` and `USER_PROMPT_S` are predefined constants used for generating the prompts.
        """
        
        system_prompt = SYSTEM_PROMPT_S
        user_prompt = USER_PROMPT_S

        body = self.load_body(client=client)
        
        return system_prompt, user_prompt.format(ps=body)
    

    def _create_model_output_scores(self) -> str:
        """
        Generates a JSON-formatted string representing the scoring output from the grading data.

        Returns:
            str: A JSON string containing scores mapped to criteria keys, e.g.,
                {"criterion2": score2, "criterion3": score3, ...}.

        Raises:
            DocumentNotLoadedError: If no document has been loaded into the instance.

        Notes:
            - The method retrieves scores using `load_grading` and converts them to integers.
            - Each score is associated with a dynamically generated key, e.g., "criterion1", "criterion2".
        """
        _, _, _, scores = self.load_grading()

        scores = [int(s) for s in scores]
        
        output = {f"factor{i+1}":s for i, s in enumerate(scores[1:])}

        return json.dumps(output)

    def _create_model_input_reasonings(self, client:Any=None, scores:str=None) -> Tuple[str, str]:
        """
        Creates input prompts for a model by combining predefined system and user prompts
        with the document's body content and scores.

        Args:
            client (Any, optional): An optional client instance for use in `load_body`. Defaults to None.
            scores (str, optional): A JSON-formatted string of scores. If not provided, scores are
                                    extracted and formatted from the grading data.

        Returns:
            Tuple[str, str]: A tuple containing:
                - system_prompt (str): The predefined system prompt.
                - user_prompt (str): The user prompt formatted with the essay body and scores.

        Notes:
            - If `scores` is not provided, it is generated by retrieving grading data using `load_grading`
            and formatting it as a JSON string.
            - `SYSTEM_PROMPT_R` and `USER_PROMPT_R` are predefined constants used for generating the prompts.
        """
        
        system_prompt = SYSTEM_PROMPT_R
        user_prompt = USER_PROMPT_R

        body = self.load_body(client=client)

        if not scores:
            _, _, _, scores = self.load_grading()
            scores = [int(s) for s in scores]
            scores = {f"factor{i+1}":score for i, score in enumerate(scores)}
            scores = json.dumps(scores)

        return system_prompt, user_prompt.format(ps=body, scores=scores)


    def _create_model_output_reasonings(self) -> str:
        """
        Generates a JSON-formatted string representing the reasonings output from the grading data.

        Returns:
            str: A JSON string containing reasonings mapped to criteria keys, e.g.,
                {"criterion1": reasoning1, "criterion2": reasoning2, ...}.

        Raises:
            DocumentNotLoadedError: If no document has been loaded into the instance.

        Notes:
            - The method retrieves reasonings using `load_grading`.
            - Each reasoning is associated with a dynamically generated key, e.g., "criterion1", "criterion2".
        """

        _, _, reasonings, _ = self.load_grading()

        output = {f"factor{i+1}":r for i, r in enumerate(reasonings)}
        output = json.dumps(output)
        
        return output
    

    def create_training_instance(self, client: Any = None, template: str = None) -> Dict[str, List[Dict[str, str]]]:
        """
        Creates a training instance for a model, formatted as a dictionary containing a sequence of messages.

        Args:
            client (Any, optional): An optional client instance used in generating model inputs. Defaults to None.
            template (str, optional): Specifies the type of template to use.
                                    - None or 's' for scores template (default).
                                    - 'r' for reasonings template.

        Returns:
            Dict[str, List[Dict[str, str]]]: A dictionary representing the training instance with
                                            system, user, and assistant messages.

        Raises:
            ValueError: If the specified `template` is not supported.
        """

        if template is None or template == 's':
            system_prompt, user_prompt = self._create_model_input_scores(client=client)
            assistant = self._create_model_output_scores()
            training_instance = {
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_prompt},
                    {'role': 'assistant', 'content': assistant},
                ]
            }

        elif template == 'r':
            system_prompt, user_prompt = self._create_model_input_reasonings(client=client)
            assistant = self._create_model_output_reasonings()
            training_instance = {
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_prompt},
                    {'role': 'assistant', 'content': assistant},
                ]
            }

        else:
            raise ValueError('Specified format is not supported.')

        return training_instance

def create_messages(
    client: Any,
    input_file: str = None,
    reader: TOKReader = None,
    template: str = None,
    examples_as_json: str = None,
    scores=None,
):
    """
    Generates a list of messages for model input based on the provided template, file, or reader.

    Args:
        client (Any): Client instance used for auxiliary tasks like generating prompts with image descriptions.
        input_file (str, optional): Path to the input file to be processed.
                                     Either `input_file` or `reader` must be specified.
        reader (TOKReader, optional): Preloaded `TOKReader` instance.
                                       Either `reader` or `input_file` must be specified.
        template (str, optional): Specifies the type of template to use.
                                   - 's' for scores.
                                   - 'r' for reasonings.
                                   If not specified, defaults to 's'.
        examples_as_json (str, optional): Path to a JSON file storing examples.
                                          New examples are created if missing and saved here.
        scores (Any, optional): Precomputed scores to include in the reasoning template.

    Returns:
        Tuple[TOKReader, List[Dict[str, str]]]:
            - `TOKReader` instance used for processing.
            - List of messages formatted for model input, including examples if provided.

    Raises:
        ValueError: If neither `input_file` nor `reader` is specified.
        ValueError: If `template` is not one of 's' or 'r'.

    Notes:
        - If `examples_as_json` is provided, examples are loaded or created and stored in the specified JSON file.
        - For scores template ('s'), the system and user prompts are generated for grading scores.
        - For reasoning template ('r'), the system and user prompts are generated for reasoning with optional scores.
    """

    new_examples_flag = False
    examples = None

    
    if examples_as_json:
        print(examples_as_json)
        if not os.path.exists(examples_as_json):
            print("path doesnt exist")
            files = []
            examples = []
        else:
            with open(examples_as_json, 'r') as f:
                print("MANAGED TO OPEN")
                files_examples = json.load(f)
                files = files_examples['files']
                examples = files_examples['examples']

        for file in glob.glob(os.path.join(EXAMPLES_PATH, '*.docx')):
            
            file_name = os.path.basename(file)
            print(file_name)
            if file_name not in files:
                example_reader = TOKReader()
                example_reader.load_doc(file)
                example = example_reader.create_training_instance(client=client, template=template)['messages'][1:]
                files.append(file_name)
                examples.extend(example)
                new_examples_flag = True

    if new_examples_flag:
        with open(examples_as_json, 'w') as f:
            json.dump({"files": files, "examples": examples}, f)

    if input_file:
        reader = TOKReader()
        reader.load_doc(input_file)
    elif reader:
        print("or here")
        pass
    else:
        raise ValueError('Either `input_file` or `reader` parameter must be specified.')

    examples = examples or []

    if template is None or template == 's':
        system_prompt, user_prompt = reader._create_model_input_scores(client)
        messages = [{'role': 'system', 'content': system_prompt}, *examples, {'role': 'user', 'content': user_prompt}]
        return reader, messages

    elif template == 'r':
        system_prompt, user_prompt = reader._create_model_input_reasonings(client, scores)
        messages = [{'role': 'system', 'content': system_prompt}, *examples, {'role': 'user', 'content': user_prompt}]
        return reader, messages
    else:
        raise ValueError(
            f'Value {template} is not recognized as allowed value. Please enter either "s" for scores or "r" for reasonings.'
        )

def catch_error(logger):
    """
    A decorator catching exceptions raised by the wrapped function, logs the error with a timestamp,
    and re-raises the exception.

    Args:
        logger (logging.Logger): Logger instance used to log error messages.

    Returns:
        function: A decorator that wraps a function to add error handling and logging.
    """

    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                logger.error(f'{datetime.now()}: {e}', exc_info=True)
                raise e

        return wrapper

    return decorator

def get_evaluation(client: Any, file, model='gpt-4o', hyperparameters=None, examples_as_json: str = None):
    """
    Evaluates TOK IB Exhibition by extracting scores and reasoning based on a set of predefined criteria,
    and generates a final grade based on the total score.

    Args:
        client (Any): The client instance used to interact with the model for generating scores and reasonings.
        file: The input file (PDF or DOCX format) containing the essay to evaluate.
        model (str): The model to use for evaluation (default is 'gpt-4o-mini').
        hyperparameters (dict, optional): A dictionary of hyperparameters to customize the model's behavior (default is None).
        examples_json (str, optional): Path to a JSON file containing example scores and reasonings (default is None).

    Returns:
        str: A formatted string containing the final grade, individual criterion scores, and the associated reasoning.
    """
    hyperparameters = hyperparameters or {}

    # examples_as_json_scores = 'scores_' + examples_as_json if examples_as_json else None
    # examples_as_json_reasonings = 'reasonings_' + examples_as_json if examples_as_json else None

    examples_as_json_scores = os.path.join(EXAMPLES_PATH, 'scores_examples.json') if examples_as_json else None
    examples_as_json_reasonings = os.path.join(EXAMPLES_PATH, 'reasonings_examples.json') if examples_as_json else None

    # SCORES
    reader, scores_messages = create_messages(client=client, input_file=file, examples_as_json=examples_as_json_scores)

    # Get scores
    scores = client.beta.chat.completions.parse(
        messages=scores_messages, model=model, **hyperparameters, response_format=ScoresExtraction
    )

    scores_json = json.loads(scores.choices[0].message.content)

    # Get grade:
    grade = np.array([v for _, v in scores_json.items()]).sum()

    # REASONINGS
    _, reasonings_messages = create_messages(
        client, reader=reader, template='r', examples_as_json=examples_as_json_reasonings, scores=scores_json
    )

    reasonings = client.beta.chat.completions.parse(
        messages=reasonings_messages, model=model, response_format=ReasoningsExtraction
    )

    reasonings_json = json.loads(reasonings.choices[0].message.content)

    output = f'''
    Grade: {grade}/35
    
    1. Relevance and Motivation.
    Score: {scores_json['factor1']}/10
    Why? {reasonings_json['factor1']}

    2. Experience and Achievements.
    Score: {scores_json['factor2']}/10
    Why? {reasonings_json['factor2']}
    
    3. Skills and Personal Qualities.
    Score: {scores_json['factor3']}/5
    Why? {reasonings_json['factor3']}
    
    4. Structure and Coherence.
    Score: {scores_json['factor4']}/5
    Why? {reasonings_json['factor4']}
    
    5. Language and Presentation.
    Score: {scores_json['factor5']}/5
    Why? {reasonings_json['factor5']}
    
    '''
    return output
    