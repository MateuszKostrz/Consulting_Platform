#!/usr/bin/env python
"""
Module Name: gpt_tools
"""

from docx import Document
from typing import Tuple, List, Dict, Any, Union
import glob
import os
import re
import json
import numpy as np
from datetime import datetime
import base64
from pdf2docx import Converter
from pydantic import BaseModel
from functools import wraps
import textwrap
import io
from io import TextIOWrapper
from werkzeug.datastructures.file_storage import FileStorage
from django.core.files.uploadedfile import InMemoryUploadedFile
from decouple import config

# CONFIG

SYSTEM_PROMPT_L = '''
Act as an expert Theory of Knowledge (TOK) examiner. 

Use the following step-by-step instructions to respond to user input.
1. **Essay Submission**: User will submit an essay enclosed in `<essay>` XML tags.

2. **Evaluation**: Evaluate the essay based on the criteria provided within `<criteria>` XML tags. 
- Consider how well every criterion is met and assign score using ***Scoring Guidelines*** below.

3. **Scoring Instructions**:
- For each criterion, use the following evaluation scale to determine how well the criterion is met.
    - **Criterion 1 - scale**: 0-5
    - **Criterion 2 - scale**: 0-10
    - **Criterion 3 - scale**: 0-10
    - **Criterion 4 - scale**: 0-10
    - **Criterion 5 - scale**: 0-10
    - **Criterion 6 - scale**: 0-5

    **Scoring Guidelines**:
    - Criterion 1: 
        - Assign 4-5 points if: The essay clearly addresses the title and maintains a focused discussion on the research question throughout the peace. It includes only information necessary to critically explore the title which is clearly reflected in the body of the essay.
        - Assign 2-3 points if: The essay adequately addresses the title and maintains a focused discussion without significant digression. Information chosen is largely relevant to the title throughout the essay.
        - Assign 0-1 points if: The title is not clearly stated, and its focus is not clear throughout the essay. There is significant digression from gist and irrelevant information is often included.
    - Criterion 2:
        - Assign 9-10 points if: The essay demonstrates a strong and clear connection between multiple Areas of Knowledge (AOKs). Each AOK is explored in depth, and the relationships between them are thoughtfully analyzed in relation to the prescribed title.
        - Assign 6-8 points if: The essay links multiple AOKs, but the connections are not fully developed or consistently clear. There is an adequate exploration of each AOK, but the relationships between them could be more explicitly analyzed.
        - Assign 3-5 points if: The essay addresses different AOKs but fails to explore the relationships between them. Connections may be superficial, and the exploration of each AOK lacks depth.
        - Assign 0-2 points if: The essay does not link different AOKs effectively, or it only focuses on one AOK. There is little to no analysis of how the AOKs relate to the title.
    - Criterion 3:
        - Assign 9-10 points if: The examples used in the essay are highly relevant and convincingly support the arguments. They are well-explained and directly tied to the discussion, helping to advance the analysis
        - Assign 6-8 points if: The examples are generally relevant and support the argument, but some may be less effective in illustrating key points. There may be occasional gaps in the explanation or linkage to the main argument.
        - Assign 3-5 points if: Examples are provided but are either not always relevant or not effectively explained. The connection between the examples and the argument may be weak or unclear.
        - Assign 0-2 points if: Few or no relevant examples are used, and the examples that are provided fail to effectively support the argument.
    - Criterion 4:
        - Assign 9-10 points if: The essay critically evaluates different perspectives and counterpoints, thoughtfully considering alternative viewpoints and their implications. The  argument is nuanced and incorporates these considerations into a well-rounded discussion.
        - Assign 6-8 points if: The essay considers different points of view and some counterpoints, but the evaluation may lack depth. While alternative perspectives are addressed, they are not fully integrated into the overall argument.
        - Assign 3-5 points if: Limited consideration is given to alternative perspectives or counterpoints. The essay may mention other viewpoints but does not evaluate them in a meaningful way.
        - Assign 0-2 points if: The essay fails to consider alternative perspectives or counterpoints. The argument is one-sided and lacks critical engagement with different points of view.
    - Criterion 5:
        - Assign 9-10 points if: The essay demonstrates a strong critical approach, with arguments that are well-developed and supported by evidence. The student consistently engages with the material in a thoughtful, analytical manner, avoiding unnecessary description.
        - Assign 6-8 points if: The essay is generally critical and argumentative but may include some descriptive sections. There is a good balance between argumentation and description, though more critical analysis could be included.
        - Assign 3-5 points if: The essay is more descriptive than argumentative. While there is some analysis, much of the content is focused on explaining rather than critically engaging with the topic
        - Assign 0-2 points if: The essay is largely descriptive and lacks critical analysis. There is little to no argumentation, and the content is mostly focused on explaining concepts rather than engaging with them critically.
    - Criterion 6:
        - Assign 4-5 points if: The essay is well-organized, with a clear introduction that sets up the argument, well-structured body paragraphs that develop the points, and a conclusion that effectively summarizes the discussion. The structure enhances the clarity of the argument. The work is well referenced both in the bibliography and in-text citation.
        - Assign 2-3 points if: The essay has a generally clear structure, but there may be issues with the flow between sections or the organization within paragraphs. The introduction, body, and conclusion are present but could be more clearly defined or better connected. Bibliography given with some mistakes or no in-text citations given.
        - Assign 0-1 points if: The essay lacks clear organization. The introduction, body, and conclusion may be poorly defined, and the argument is difficult to follow due to structural issues. There may be significant jumps between ideas, making the essay hard to read. No bibliography given.

4. **Marking**: For each criterion assign a score reflecting how well the essay meets the criterion. 
Use ***Scoring Guidelines*** when in doubt how to evaluate a particular criterion.
**Be strict** but fair — only award high scores when the essay clearly meets the expectations.

4. **Output**: Provide the results in a JSON format as follows: {{"criterion1":3, "criterion2":7, ..., "criterion6":4}}

**Evaluation Criteria**:
<criteria>
1. The essay's title is clearly stated and the essay's focus is sustained on the title chosen and does not include digression or irrelevant information.
2. The work is effectively linked with  different Areas of Knowledge.
3. The students provides examples which are convincing and support the argument being made efficiently.
4. Any implications, counterpoints of the arguments, and different points of views  are considered and evaluated  by the student.
5. The essay is predominately \ncritical and argumentative rather than descriptive.
6. The work is well-organized with an introductory paragraph, main body sections, and a concluding paragraph.







</criteria>
'''

USER_PROMPT_L = '''
<essay>
{essay}
</essay>'''


SYSTEM_PROMPT_R = '''
Act as an expert Theory of Thought (TOK) examiner. 
Use the following step-by-step instructions to respond to user input.

1. **Submission**: User will submit an essay enclosed in `<essay>` XML tags and a list of scores enclosed in `<scores>` XML tags. 
- Scores indicate how well criteria (listed within `<criteria>` XML tags) are met based on the ***Scoring Guidelines*** below:

**Scoring Guidelines**:
- Criterion 1 (0-5 points): 
    - 4-5 points: The essay clearly addresses the title and maintains a focused discussion on the research question throughout the peace. It includes only information necessary to critically explore the title which is clearly reflected in the body of the essay.
    - 2-3: The essay adequately addresses the title and maintains a focused discussion without significant digression. Information chosen is largely relevant to the title throughout the essay.
    - 0-1 points: The title is not clearly stated, and its focus is not clear throughout the essay. There is significant digression from gist and irrelevant information is often included.
- Criterion 2 (0-10 points):
    - 9-10 points: The essay demonstrates a strong and clear connection between multiple Areas of Knowledge (AOKs). Each AOK is explored in depth, and the relationships between them are thoughtfully analyzed in relation to the prescribed title.
    - 6-8 points: The essay links multiple AOKs, but the connections are not fully developed or consistently clear. There is an adequate exploration of each AOK, but the relationships between them could be more explicitly analyzed.
    - 3-5 points: The essay addresses different AOKs but fails to explore the relationships between them. Connections may be superficial, and the exploration of each AOK lacks depth.
    - 0-2 points: The essay does not link different AOKs effectively, or it only focuses on one AOK. There is little to no analysis of how the AOKs relate to the title.
- Criterion 3 (0-10 points):
    - 9-10 points: The examples used in the essay are highly relevant and convincingly support the arguments. They are well-explained and directly tied to the discussion, helping to advance the analysis
    - 6-8 points: The examples are generally relevant and support the argument, but some may be less effective in illustrating key points. There may be occasional gaps in the explanation or linkage to the main argument.
    - 3-5 points: Examples are provided but are either not always relevant or not effectively explained. The connection between the examples and the argument may be weak or unclear.
    - 0-2 points: Few or no relevant examples are used, and the examples that are provided fail to effectively support the argument.
- Criterion 4 (0-10 points):
    - 9-10 points: The essay critically evaluates different perspectives and counterpoints, thoughtfully considering alternative viewpoints and their implications. The  argument is nuanced and incorporates these considerations into a well-rounded discussion.
    - 6-8 points: The essay considers different points of view and some counterpoints, but the evaluation may lack depth. While alternative perspectives are addressed, they are not fully integrated into the overall argument.
    - 3-5 points: Limited consideration is given to alternative perspectives or counterpoints. The essay may mention other viewpoints but does not evaluate them in a meaningful way.
    - 0-2 points: The essay fails to consider alternative perspectives or counterpoints. The argument is one-sided and lacks critical engagement with different points of view.
- Criterion 5 (0-10 points):
    - 9-10 points: The essay demonstrates a strong critical approach, with arguments that are well-developed and supported by evidence. The student consistently engages with the material in a thoughtful, analytical manner, avoiding unnecessary description.
    - 6-8 points: The essay is generally critical and argumentative but may include some descriptive sections. There is a good balance between argumentation and description, though more critical analysis could be included.
    - 3-5 points: The essay is more descriptive than argumentative. While there is some analysis, much of the content is focused on explaining rather than critically engaging with the topic
    - 0-2 points: The essay is largely descriptive and lacks critical analysis. There is little to no argumentation, and the content is mostly focused on explaining concepts rather than engaging with them critically.
- Criterion 6 (0-5 points):
    - 4-5 points if: The essay is well-organized, with a clear introduction that sets up the argument, well-structured body paragraphs that develop the points, and a conclusion that effectively summarizes the discussion. The structure enhances the clarity of the argument. The work is well referenced both in the bibliography and in-text citation.
    - 2-3 points if: The essay has a generally clear structure, but there may be issues with the flow between sections or the organization within paragraphs. The introduction, body, and conclusion are present but could be more clearly defined or better connected. Bibliography given with some mistakes or no in-text citations given.
    - 0-1 points if: The essay lacks clear organization. The introduction, body, and conclusion may be poorly defined, and the argument is difficult to follow due to structural issues. There may be significant jumps between ideas, making the essay hard to read. No bibliography given.

2. **Evaluation**: For every criterion:
   - **Interpret the Score**: Use the provided ***Scoring Guidelines*** to understand the extent to which the criterion is met.
   - **Provide Reasoning**: Offer a thorough explanation of why the criterion is met to the extent specified by the score. 
   Use specific examples from the essay to illustrate how the criterion is fulfilled. 
   Clearly indicate which parts of the essay correspond to the criterion. When appropriate cite the relevant parts of the essay.

3. **Output**: Present the reasonings in a JSON format as follows: {{"criterion1":reasoning#1, "criterion2":reasoning#2,...,"criterion6":reasoning#6}}

**Evaluation Criteria**:
<criteria>
1. The essay's title is clearly stated and the essay's focus is sustained on the title chosen and does not include digression or irrelevant information.
2. The work is effectively linked with  different Areas of Knowledge.
3. The students provides examples which are convincing and support the argument being made efficiently.
4. Any implications, counterpoints of the arguments, and different points of views  are considered and evaluated  by the student.
5. The essay is predominately \ncritical and argumentative rather than descriptive.
6. The work is well-organized with an introductory paragraph, main body sections, and a concluding paragraph.






</criteria>
'''

USER_PROMPT_R = '''
**Essay**:
<essay>
{essay}
</essay>

**Scores**:
<scores>
{labels}
</scores>
'''

TOK_TITLES = [
    'Is subjectivity overly celebrated in the arts but unfairly condemned in history? Discuss with reference to the arts and history', 
    'How can we reconcile the opposing demands for specialization and generalization in the production of knowledge? Discuss with reference to mathematics and one other area of knowledge', 
    'Nothing is more exciting than fresh ideas, so why are areas of knowledge often so slow to adopt them? Discuss with reference to the human sciences and one other area of knowledge', 
    'Do we underestimate the challenges of taking knowledge out of its original context and transferring it to a different context? Discuss with reference to two areas of knowledge', 
    'Do we need custodians of knowledge? Discuss with reference to two areas of knowledge', 
    'Are we too quick to assume that the most recent evidence is inevitably the strongest? Discuss with reference to the natural sciences and one other area of knowledge']


class ScoresExtraction(BaseModel):
    criterion1: int
    criterion2: int
    criterion3: int
    criterion4: int
    criterion5: int
    criterion6: int


class ReasoningsExtraction(BaseModel):
    criterion1: str
    criterion2: str
    criterion3: str
    criterion4: str
    criterion5: str
    criterion6: str


class DocumentNotLoadedError(Exception):
    """
    Exception raised when a document is accessed before being loaded.
    """
    def __init__(self, message="The document has not been loaded. Please load a document before accessing the method."):
        """
        Initialize the DocumentNotLoadedError with an optional error message.
        
        Args:
            message (str): The error message to be displayed.
        """
        self.message = message
        super().__init__(self.message)


def get_image_description(client: Any, b64_image: str) -> str:
    """
    Generates a textual description of the content within an image using a GPT-4-based model.

    Args:
        client (Any): An instance of an API client used to interact with the GPT-4 service.
        b64_image (str): A base64-encoded string representing the image to be described.

    Returns:
        str: A concise description of the image content as generated by the model.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Describe what's in the image. Be precise and concise."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_image}"}}
                ]
            }
        ],
        max_tokens=300,
    )
    
    return response.choices[0].message.content


class TOKReader:
    """
    Custom class to read and process TOK IB essays.
    """

    def __init__(self):
        self.namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
    
    def load_doc(self, file: Union[str, FileStorage, TextIOWrapper]) -> None:
        """
        Loads a PDF or DOCX file into a `Document` object.

        Function handles loading both PDF and DOCX formats. It converts PDF files
        into DOCX format before loading them into a `Document` object. The file can be
        provided as a string path, a Flask `FileStorage` object, or a `TextIOWrapper`.

        Args:
            file (str or FileStorage or TextIOWrapper): The file to load, either as a string
            representing the file path, or a file-like object.

        Raises:
            FileNotFoundError: If the file path is invalid or doesn't exist.
            ValueError: If the file format is unsupported (neither PDF nor DOCX).
            RuntimeError: If any other error occurs while loading the document.
        """
        try:
            print("Making it here")
            print(f"Type of file: {type(file)}")  # Debugging line

            # Handle different types of file inputs
            if isinstance(file, str):
                if not os.path.exists(file):
                    raise FileNotFoundError(f'File not found: {file}')
                _, file_ext = os.path.splitext(file)
                print(f"File extension from path: {file_ext}")  # Debugging line
                with open(file, 'rb') as f:
                    file_stream = io.BytesIO(f.read())
            elif isinstance(file, InMemoryUploadedFile):
                file_ext = os.path.splitext(file.name)[1]  # Get file extension
                print(f"File extension from InMemoryUploadedFile: {file_ext}")  # Debugging line
                file_stream = io.BytesIO(file.read())
            elif isinstance(file, TextIOWrapper):
                file_ext = os.path.splitext(file.name)[1] if hasattr(file, 'name') else ''
                print(f"File extension from TextIOWrapper: {file_ext}")  # Debugging line
                file_stream = io.BytesIO(file.read().encode('utf-8'))
            else:
                raise ValueError('Invalid file type. Must be a string path or file-like object.')

            print(f"File extension being processed: {file_ext}")  # Debugging line

            if file_ext.lower() == '.pdf':
                file_stream.seek(0)
                docx_buffer = io.BytesIO()
                converter = Converter(stream=file_stream)
                converter.convert(docx_buffer)
                docx_buffer.seek(0)
                self.doc = Document(docx_buffer)
            elif file_ext.lower() == '.docx':
                self.doc = Document(file_stream)
            else:
                raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")
        except Exception as e:
            print(f"[ERROR] Error processing uploaded file: {e}")  # Updated error message
            raise RuntimeError(f"An error occurred while loading the document: {e}")






    def load_body(self, client=None) -> str:
        """
        Processes the document loaded with .load_doc() method.

        Method reads the document extracting text and image descriptions.
        It accumulates paragraphs until it encounters a paragraph starting with the 'grade' or 'grading criteria'
        or 'total points' string., 
        and processes any embedded images by appending a description of each image.

        If the document content has been previously cached, it returns the cached content. Otherwise,
        it processes the document, extracts the relevant text and image descriptions, and caches the result.

        Args:

        Returns:
            str: Body of the document.

        Raises:
            DocumentNotLoadedError: If the document has not been loaded properly before calling this method.
        """

        if not hasattr(self, 'doc'):
            raise DocumentNotLoadedError()
            

        if hasattr(self, '_body_cache') and self._body_cache is not None:
            return self._body_cache

        paragraphs = []

        if not self.doc.element.body.findall('.//w:drawing', self.namespaces):
            for paragraph in self.doc.paragraphs:
                text = paragraph.text.strip()
                if text.lower().startswith(('grade', 'grading criteria', 'total points')):
                    break
                paragraphs.append(text)

            self._body_cache = " ".join(paragraphs)
            return self._body_cache
        
        for paragraph in self.doc.element.body.findall('.//w:p', self.namespaces):
            text = paragraph.text.strip()
            if text.lower().startswith(('grade', 'grading criteria', 'total points')):
                break
            
            paragraphs.append(text)

            drawings = paragraph.findall('.//w:drawing', self.namespaces)

            if drawings:
                for drawing in drawings: 
                    paragraphs.append('\n[IMAGE DESCRIPTION START]\n')
                    
                    blip = drawing.find('.//a:blip', namespaces=self.namespaces)
                    image_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    image = self.doc.part.related_parts[image_id].blob
                    b64_image = base64.b64encode(image).decode('utf-8')

                    # Sends image to OpenAI.
                    image_description = get_image_description(client, b64_image)

                    paragraphs.append(image_description)

                    paragraphs.append('\n[IMAGE DESCRIPTION END]\n')
                    
        return ' '.join(paragraphs)

    
    def load_grading(self) -> Tuple[List[str], List[str], List[str]]:
        """
        Extracts and loads grading criteria, scales, reasonings, and labels from the document.

        Method searches through tables in the document to extract grading information. It looks 
        for tables and iterates through the rows to extract data from each cell. The expected 
        format is four columns per row: criteria, scales, reasonings, and labels. The labels should 
        be digits. The method caches the results for future calls to improve performance.

        Returns:
            Tuple[List[str], List[str], List[str], List[str]]:
                - A list of grading criteria.
                - A list of scales corresponding to each criterion.
                - A list of reasonings associated with each criterion.
                - A list of scores for each criterion, expected to be digits.

        Raises:
            DocumentNotLoadedError:
                If the document has not been loaded before this method is called.
            
            ValueError:
                If no tables are found in the document.

        Notes:
            - The method expects the document contains tables structured with exactly four columns
            per row. 
            - Rows must have data in all four columns. Grade must be a digit. Otherwise the row is skipped.
            - The method caches the extracted data in `_grading_cache` to avoid redundant processing in
            subsequent calls.
        """

        if not hasattr(self, 'doc'):
            raise DocumentNotLoadedError()
        
        if not self.doc.element.body.findall('.//w:tbl', self.namespaces):
            raise ValueError('Table not found.')

        if hasattr(self, '_grading_cache') and self._grading_cache is not None:
            return self._grading_cache
            
        criteria = []
        scales = []
        reasonings = []
        labels = []

        for table in self.doc.element.body.findall('.//w:tbl', self.namespaces):
            for row in table.findall('w:tr', self.namespaces):
                c, s, r, g = [' '.join([p.text for p in cell.findall('w:p', self.namespaces)]).strip() for cell in row.findall('w:tc', self.namespaces)]
                if not (c and s and r and g and g.isdigit()):
                    continue
                criteria.append(c)
                scales.append(s)
                reasonings.append(r)
                labels.append(g)

        self._grading_cache = (criteria, scales, reasonings, labels)
        return self._grading_cache
        
    
    def create_model_input_labels(self, client:Any=None) -> Tuple[str, str]:
        """
        Creates formatted system and user messasge for scores generation task.

        Method generates two strings, `system_prompt` and `user_prompt`, formatted with the 
        the essay content extracted from the document.

        Args:
            client (Any, optional): An optional client (client is only neccessary 
            for image-containing TOKs as images are sent to OpenAI for processing).
            Defaults to None.

        Returns:
            Tuple[str, str]:
                - A formatted `system_prompt` string that includes the grading criteria and scales.
                - A formatted `user_prompt` string that includes the essay content.

        Notes:
            - The `SYSTEM_PROMPT_L` and `USER_PROMPT_L` constants should be predefined templates 
            for formatting the prompts.
            - The method relies on the `load_body` and `load_grading` methods to retrieve the essay 
            and grading information, respectively.
        """
        
        system_prompt = SYSTEM_PROMPT_L
        user_prompt = USER_PROMPT_L

        essay = self.load_body(client=client)
        
        return system_prompt, user_prompt.format(essay=essay)
    

    def create_model_output_labels(self) -> str:
        """
        Creates a JSON-formatted string of scores.

        Method retrieves the grading labels from the document using the `load_grading` method, 
        and then converts these labels into a JSON-formatted string used as AI message in scores
        generation task.

        Returns:
            str: A JSON-formatted string containing the grading labels.

        Notes:
            - The `load_grading` method is expected to return four lists, where the fourth list 
            contains the grading labels.
            - The labels are converted to a JSON string using `json.dumps`.
        """
        _, _, _, labels = self.load_grading()

        labels = [int(l) for l in labels]
        
        output = {f"criterion{i+1}":l for i, l in enumerate(labels)}

        return json.dumps(output)


    def create_model_input_reasonings(self, client:Any=None, labels:str=None) -> Tuple[str, str]:
        """
        Creates formatted system and user messasge for reasonings generation task.

        Method generates two strings, `system_prompt` and `user_prompt`, formatted with the 
        essay content. It allows the option to use provided labels (`labels`) 
        instead of the default labels retrieved from the document. 

        Args:
            client (Any, optional): An optional client object used for processing, if needed. Defaults to None.
            labels (str, optional): A JSON-formatted string of scores dictionary. If not provided, the labels are fetched from the document.

        Returns:
            Tuple[str, str]:
                - A formatted `system_prompt` string that includes the grading criteria and scales.
                - A formatted `user_prompt` string that includes the essay content and reasonings/labels.

        Notes:
            - The `SYSTEM_PROMPT_R` and `USER_PROMPT_R` constants should be predefined templates 
            for formatting the prompts.
            - The method relies on the `load_body` and `load_grading` methods to retrieve the essay 
            and grading information, respectively.
            - If `labels` is not provided, the method uses the default labels from the document.
        """
        
        system_prompt = SYSTEM_PROMPT_R
        user_prompt = USER_PROMPT_R

        essay = self.load_body(client=client)

        if not labels:
            _, _, _, labels = self.load_grading()
            labels = [int(l) for l in labels]
            labels = {f"criterion{i+1}":label for i, label in enumerate(labels)}
            labels = json.dumps(labels)

        return system_prompt, user_prompt.format(essay=essay, labels=labels)


    def create_model_output_reasonings(self) -> str:
        """
        Creates a JSON-formatted string of grading reasonings.

        Method retrieves the grading reasonings from the document using the `load_grading` method,
        and then converts the reasonings into a JSON-formatted string.

        Returns:
            str: A JSON-formatted string containing the grading reasonings.

        Notes:
            - The `load_grading` method is expected to return four lists, where the third list contains 
            the grading reasonings.
            - The reasonings are converted to a JSON string using `json.dumps`.
        """

        _, _, reasonings, _ = self.load_grading()

        output = {f"criterion{i+1}":r for i, r in enumerate(reasonings)}
        output = json.dumps(output)
        
        return output
    

    def create_training_instance(self, client:Any, template:str) -> Dict[str, List[Dict[str, str]]]:
        """
        Creates a training instance for the GPT model based on the provided template type.

        Depending on the template type, the function constructs a structured input for the model 
        containing system, user, and assistant prompts. It supports two template formats:
        'l' for label-(or scores)-generation task and 'r' for reasoning-generation task.
        Each instance is returned as a dictionary with a list of message components, where each message includes 
        a role ('system', 'user', 'assistant') and its corresponding content.

        Args:
            client (Any): The GPT client used to generate prompts for the model.
            template (str): The template type, either 'l' for label-generation task 
            or 'r' for reasoning-generation task.

        Returns:
            Dict[str, List[Dict[str, str]]]: A dictionary representing the training instance, where
            the key 'messages' contains a list of dictionaries, each with 'role' and 'content'.

        Raises:
            ValueError: If the provided template is neither 'l' nor 'r'.
        """
        if template == 'l':
            system_prompt, user_prompt = self.create_model_input_labels(client=client)
            assistant = self.create_model_output_labels()
            training_instance = {'messages':
                                [{'role':'system', 'content':system_prompt}, 
                                 {'role':'user', 'content': user_prompt}, 
                                 {'role':'assistant', 'content': assistant}]}
            
        elif template == 'r':
            system_prompt, user_prompt = self.create_model_input_reasonings(client=client)
            assistant = self.create_model_output_reasonings()
            training_instance = {'messages': 
                                 [{'role':'system', 'content':system_prompt},
                                  {'role':'user', 'content':user_prompt},
                                  {'role':'assistant', 'content':assistant}]}
            
        else:
            raise ValueError('Specified format is not supported.')
        
        return training_instance


def catch_error(logger):
    """
    A decorator wrapping a function with error handling and logging.

    Decorator logs any exception raised during the execution of the wrapped function.
    It catches all exceptions, logs the error with a timestamp, and returns a generic error 
    message along with an HTTP 500 status code if an error occurs.

    Args:
        logger (logging.Logger): A logger instance used to log errors.

    Returns:
        function: A wrapped function that includes error handling.

    Example:
        @catch_error(logger)
        def some_function():
            # Function logic
    """
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                logger.error(f'{datetime.now()}. Error in {func.__name__}: {e}', exc_info=True)
                return 'An error occured during processing', 500
        return wrapper
    return decorator


def get_evaluation(client:Any, 
                   file, 
                   model='gpt-4o-mini', 
                   hyperparameters=None, 
                   examples_dir:str=None, 
                   tok_titles:list=None):
    """
    Generates an evaluation for a Theory of Knowledge (TOK) essay based on model predictions.

    Function extracts content from DOCX or PDF document, and generates 
    evaluation using a GPT-4o-mini model. The function returns a formatted string.

    Args:
        client (Any): The OpenAI client used to interact with the model for score and reasoning generation.
        file (str or FileStorage): The essay document to evaluate.
        model (str, optional): The model name to use for GPT predictions. Defaults to 'gpt-4o-mini'.
        hyperparameters (dict, optional): Hyperparameters for the GPT model (e.g., temperature, top_p). 
                                          Defaults to None.
        examples_dir (str, optional): Path to a directory containing documents for generating examples for the model.
            Defaults to None.
        tok_titles (list, optional): A list of predefined TOK titles for title extraction. Defaults to None.

    Returns:
        str: A formatted string containing the evaluation report, including the essay's title, grade, 
             individual scores for each criterion, and the reasoning behind those scores.

    Raises:
        ValueError: If the specified file or template type is invalid.
    """
    # Configure hyperparameters.
    hyperparameters = hyperparameters or {}

    examples_l = None
    examples_r = None

    # Get examples.
    if examples_dir is not None:
        examples_r=[]
        examples_l=[]
        for example in glob.glob(os.path.join(examples_dir,'*.docx')):
            reader = TOKReader()
            reader.load_doc(example)
            example_l = reader.create_training_instance(client=client, template='l')['messages'][1:]
            example_r = reader.create_training_instance(client=client, template='r')['messages'][1:]
            examples_l.extend(example_l)
            examples_r.extend(example_r)

    examples_l = examples_l or []
    examples_r = examples_r or []

    reader = TOKReader()
    reader.load_doc(file=file)

    # Create messages for scores generation.
    system_prompt_l, user_prompt_l = reader.create_model_input_labels(client=client)
    messages_l = [{'role':'system', 'content': system_prompt_l},
                *examples_l, 
                {'role':'user', 'content': user_prompt_l}]

    # Get scores.
    scores = client.beta.chat.completions.parse(
        messages=messages_l, 
        model=model,
        **hyperparameters,
        response_format=ScoresExtraction)

    # Extract scores dictionary:
    scores_json = json.loads(scores.choices[0].message.content)

    # Extract grade:
    scores_sum = np.array([v for _, v in scores_json.items()]).sum()
    if scores_sum >= 45:
        grade =  'A'
    elif scores_sum >= 35:
        grade =  'B'
    elif scores_sum >= 25:
        grade = 'C'
    elif scores_sum >= 15:
        grade =  'D'
    else:
        grade =  'E'

    # Create messages for reasonings generations:
    system_prompt_r, user_prompt_r = reader.create_model_input_reasonings(client, scores_json)
    messages_r = [{'role':'system', 'content': system_prompt_r},
                *examples_r,
                {'role':'user', 'content': user_prompt_r}]

    # Get reasonings:
    reasonings = client.beta.chat.completions.parse(
        messages=messages_r, 
        model=model, 
        response_format=ReasoningsExtraction
    )

    reasonings_json = json.loads(reasonings.choices[0].message.content)


    PDF_TEMPLATE = f'''
Grade: {grade}
    
1. The essay's title is clearly stated and the essay's focus is sustained on the title chosen and
does not include digression or irrelevant information.
Score: {scores_json['criterion1']} / 5

Reasoning: {textwrap.fill(reasonings_json['criterion1'], width=80)}

    
2. The work is effectively linked with  different Areas of Knowledge.
Score: {scores_json['criterion2']} / 10

Reasoning: {textwrap.fill(reasonings_json['criterion2'], width=80)}
    

3. The students provides examples which are convincing and support  the argument being made
efficiently.

Score: {scores_json['criterion3']} / 10

Reasoning: {textwrap.fill(reasonings_json['criterion3'], width=80)}
    

4. Any implications, counterpoints of the arguments, and different  points of views are considered
and evaluated  by the student.

Score: {scores_json['criterion4']} / 10

Reasoning: {textwrap.fill(reasonings_json['criterion4'], width=80)}
    

5. The essay is predominately \ncritical and argumentative rather than  descriptive.

Score: {scores_json['criterion5']} / 10

Reasoning: {textwrap.fill(reasonings_json['criterion5'], width=80)}
    

6. The work is well-organized with an introductory paragraph, main body sections, and a concluding 
paragraph.

Score: {scores_json['criterion6']} / 5

Reasoning: {textwrap.fill(reasonings_json['criterion6'], width=80)}
    '''

    output = PDF_TEMPLATE
    return output
    
